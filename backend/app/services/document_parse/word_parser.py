"""Word parser using python-docx.

Extracts text paragraph by paragraph from Word documents and returns
structured data with paragraph indices and metadata.
"""

import logging
from typing import List, Dict, Any

from docx import Document

logger = logging.getLogger("course_agent.document_parse.word")


def parse_word(file_path: str) -> List[Dict[str, Any]]:
    """Extract text from a Word document paragraph by paragraph.

    Args:
        file_path: Absolute or relative path to the .doc or .docx file.

    Returns:
        A list of dicts, each containing:
            - paragraph_index (int): 1-based paragraph number.
            - content (str): Text of the paragraph.
            - metadata (dict): File-level metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
    """
    results: List[Dict[str, Any]] = []

    doc = Document(file_path)
    total_paragraphs = len(doc.paragraphs)
    logger.info(f"Parsing Word document: {file_path}, {total_paragraphs} paragraphs")

    # Collect document-level metadata
    doc_metadata = {
        "total_paragraphs": total_paragraphs,
        "file_path": file_path,
    }

    # Try to extract core properties
    try:
        core_props = doc.core_properties
        if core_props.title:
            doc_metadata["title"] = core_props.title
        if core_props.author:
            doc_metadata["author"] = core_props.author
        if core_props.created:
            doc_metadata["created"] = str(core_props.created)
    except Exception:
        pass

    for para_index, para in enumerate(doc.paragraphs, start=1):
        try:
            text = para.text.strip()
            if text:
                results.append({
                    "paragraph_index": para_index,
                    "content": text,
                    "metadata": {
                        **doc_metadata,
                        "style": para.style.name if para.style else None,
                    },
                })
        except Exception as exc:
            logger.warning(f"Failed to extract paragraph {para_index}: {exc}")

    # Also extract text from tables
    table_para_index = total_paragraphs + 1
    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells_text.append(cell_text)
            if cells_text:
                content = " | ".join(cells_text)
                results.append({
                    "paragraph_index": table_para_index,
                    "content": content,
                    "metadata": {
                        **doc_metadata,
                        "style": "table",
                        "table_index": table_index,
                        "row_index": row_index,
                    },
                })
                table_para_index += 1

    return results
